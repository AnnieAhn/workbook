"""
워크북 Word(.docx) 파일 생성 모듈
claude_generator.py 에서 받은 content dict → bytes (다운로드용)
"""
import io
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 색상 ──────────────────────────────────────────────
C_DARK   = RGBColor(0x1F, 0x4E, 0x79)
C_MID    = RGBColor(0x2E, 0x74, 0xB5)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_GRAY   = RGBColor(0x60, 0x60, 0x60)

BG_SECTION = "D6E4F0"
BG_LIGHT   = "EBF3FB"
BG_YELLOW  = "FFF2CC"
BG_GREEN   = "E2EFDA"
BG_GRAY    = "F2F2F2"


# ── 유틸 ──────────────────────────────────────────────
def _shd(elem, hex_color: str):
    pPr = elem.get_or_add_pPr() if hasattr(elem, 'get_or_add_pPr') else elem
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def set_para_bg(para, hex_color):
    _shd(para._p.get_or_add_pPr(), hex_color)

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def p(doc, sb=4, sa=3, indent=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    return para

def run(para, text, bold=False, size=10, color=None, ul=False, italic=False):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic; r.underline = ul
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r

def section_title(doc, text):
    para = p(doc, 10, 4)
    set_para_bg(para, BG_SECTION)
    run(para, f"  {text}", bold=True, size=13, color=C_DARK)

def sub_title(doc, text):
    para = p(doc, 7, 3)
    run(para, text, bold=True, size=11, color=C_MID)

def blue_box(doc, text):
    para = p(doc, 3, 3, indent=0.3)
    para.paragraph_format.right_indent = Cm(0.3)
    set_para_bg(para, BG_LIGHT)
    run(para, text, size=9.5)

def col_w(table, idx, w_cm):
    for row in table.rows:
        row.cells[idx].width = Cm(w_cm)

def blank_row(doc):
    doc.add_paragraph()

# ── 선택지 출력 ───────────────────────────────────────
def add_choices(doc, choices, indent=0.8):
    para = p(doc, 2, 5, indent=indent)
    for ch in choices:
        run(para, f"  {ch}  ", size=10)

# ── 정답 테이블 ───────────────────────────────────────
def answer_table(doc, rows_data, col_headers, col_widths, bg=BG_GREEN):
    tbl = doc.add_table(rows=1, cols=len(col_headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(col_headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
    for row_data in rows_data:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    for i, w in enumerate(col_widths):
        col_w(tbl, i, w)
    return tbl


# ═══════════════════════════════════════════════════════
# 메인 빌더
# ═══════════════════════════════════════════════════════
def build_workbook_docx(
    content: dict,
    passage: str,
    passage_no: str = "",
    topic_hint: str = "",
    include_grammar: bool = True,
    include_essay: bool = True,
) -> bytes:

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
    doc.styles['Normal'].font.name = '맑은 고딕'
    doc.styles['Normal'].font.size = Pt(10)

    analysis   = content.get("analysis", {})
    vocab      = content.get("vocabulary", [])
    tf_qs      = content.get("tf_questions", [])
    today      = datetime.date.today().strftime("%Y년 %m월 %d일")
    topic      = topic_hint or analysis.get("topic_kor", "영어 지문")

    # ────────────────────────────────────────────────────
    # 표지
    # ────────────────────────────────────────────────────
    cover = p(doc, 30, 5)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(cover, "📘  완전정복 워크북", bold=True, size=18, color=C_DARK)

    c2 = p(doc, 5, 5)
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(c2, topic, bold=True, size=14, color=C_MID)

    if passage_no:
        c3 = p(doc, 3, 3)
        c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(c3, passage_no, size=10, color=C_GRAY)

    c4 = p(doc, 20, 3)
    c4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(c4, f"날짜: {today}          이름: ___________________          점수: _______ / 100점", size=11)
    doc.add_page_break()

    # ────────────────────────────────────────────────────
    # SECTION 1: 지문 읽기
    # ────────────────────────────────────────────────────
    section_title(doc, "📖  SECTION 1 : 지문 읽기")
    blue_box(doc, "⏱ 처음에는 사전/번역 없이 2분 안에 읽어보세요. 모르는 단어는 동그라미만 표시하고 계속 읽으세요.")

    psg = p(doc, 6, 6)
    psg.paragraph_format.line_spacing = Pt(22)
    psg.paragraph_format.left_indent  = Cm(0.3)
    psg.paragraph_format.right_indent = Cm(0.3)
    run(psg, passage, size=11)
    blank_row(doc)

    # ────────────────────────────────────────────────────
    # SECTION 2: 핵심 어휘
    # ────────────────────────────────────────────────────
    doc.add_page_break()
    section_title(doc, f"📝  SECTION 2 : 핵심 어휘 정복 ({len(vocab)}개)")
    blue_box(doc, "아래 표를 먼저 학습한 후 → 어휘 확인 테스트에 도전하세요!")

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['No', '단어/표현', '품사', '한국어 뜻', '지문 속 예문', '동의어']):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, BG_GRAY)
    for v in vocab:
        row = tbl.add_row()
        vals = [v.get('no',''), v.get('word',''), v.get('pos',''),
                v.get('meaning',''), v.get('example',''), v.get('synonym','')]
        for i, val in enumerate(vals):
            row.cells[i].text = str(val)
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)
            if i <= 2:
                row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 1:
                row.cells[i].paragraphs[0].runs[0].bold = True
    for i, w in enumerate([0.7, 2.8, 1.1, 2.8, 5.5, 2.6]):
        col_w(tbl, i, w)

    # 어휘 확인 테스트
    blank_row(doc)
    sub_title(doc, "✏️  어휘 확인 테스트  —  한국어 뜻을 보고 영어 단어/표현을 쓰세요")
    test_words = content.get("vocab_test_words", vocab[:10])
    # Word Box
    word_box_words = [t.get('answer', t.get('word','')) for t in test_words[:10]]
    blue_box(doc, "📦 Word Box:  " + "  /  ".join(word_box_words))

    vt = doc.add_table(rows=len(test_words[:10])//2 + len(test_words[:10])%2 + 1, cols=2)
    vt.style = 'Table Grid'
    set_cell_bg(vt.rows[0].cells[0], BG_GRAY)
    set_cell_bg(vt.rows[0].cells[1], BG_GRAY)
    vt.rows[0].cells[0].text = "한국어 뜻  →  영어 쓰기"
    vt.rows[0].cells[1].text = "한국어 뜻  →  영어 쓰기"
    for cell in vt.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for idx, t in enumerate(test_words[:10]):
        row_idx = idx // 2 + 1
        col_idx = idx % 2
        kor = t.get('meaning_kor', t.get('meaning', ''))
        vt.rows[row_idx].cells[col_idx].text = f"{idx+1}. {kor}  →  {'_'*18}"
        vt.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9.5)
    col_w(vt, 0, 7.3); col_w(vt, 1, 7.3)

    # ────────────────────────────────────────────────────
    # SECTION 3: 구조 분석
    # ────────────────────────────────────────────────────
    doc.add_page_break()
    section_title(doc, "🔍  SECTION 3 : 지문 구조 분석")

    sub_title(doc, "3-1.  문단별 핵심 요약  (한국어로 직접 써보세요)")
    summaries = analysis.get("paragraph_summaries", [])
    labels_kr = ["[도입]", "[전개]", "[심화]", "[결론]", "[보충]"]
    st_tbl = doc.add_table(rows=max(len(summaries),2), cols=2)
    st_tbl.style = 'Table Grid'
    for i in range(max(len(summaries),2)):
        label = f"문단 {i+1}\n{labels_kr[i] if i < len(labels_kr) else ''}"
        st_tbl.rows[i].cells[0].text = label
        st_tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        st_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        st_tbl.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(st_tbl.rows[i].cells[0], BG_LIGHT)
        st_tbl.rows[i].cells[1].paragraphs[0].add_run("\n\n").font.size = Pt(10)
    col_w(st_tbl, 0, 2.5); col_w(st_tbl, 1, 12.0)

    blank_row(doc)
    sub_title(doc, "3-2.  논리 흐름 도식")
    flow_p = p(doc, 5, 5)
    flow_p.paragraph_format.left_indent = Cm(0.5)
    flow_p.paragraph_format.line_spacing = Pt(22)
    set_para_bg(flow_p, BG_LIGHT)
    struct_type = analysis.get("structure_type", "")
    sum_texts = analysis.get("paragraph_summaries", ["도입","전개","결론"])
    flow_parts = []
    for i, s in enumerate(sum_texts[:4]):
        short = s[:20] + "..." if len(s) > 20 else s
        flow_parts.append(f"[문단{i+1}] {short}")
    run(flow_p, f"  구조: {struct_type}\n\n  " + "\n     ↓\n  ".join(flow_parts), size=10)

    blank_row(doc)
    sub_title(doc, "3-3.  핵심 연결어 / 핵심 표현 분석")
    connectors = analysis.get("key_connectors", [])
    if connectors:
        conn_tbl = doc.add_table(rows=1, cols=3)
        conn_tbl.style = 'Table Grid'
        for i, h in enumerate(['연결어 / 핵심 표현', '위치', '기능 / 역할']):
            cell = conn_tbl.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(cell, BG_GRAY)
        for c in connectors:
            row = conn_tbl.add_row()
            row.cells[0].text = c.get('expression', '')
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[0].paragraphs[0].runs[0].italic = True
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            row.cells[1].text = c.get('location', '')
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
            row.cells[2].text = c.get('function', '')
            row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
        col_w(conn_tbl, 0, 3.5); col_w(conn_tbl, 1, 3.5); col_w(conn_tbl, 2, 7.5)

    # ────────────────────────────────────────────────────
    # SECTION 4: 기본 독해 확인
    # ────────────────────────────────────────────────────
    doc.add_page_break()
    section_title(doc, "✅  SECTION 4 : 기본 독해 확인")

    sub_title(doc, "4-1.  T / F 문제  —  맞으면 T, 틀리면 F를 쓰고, 반드시 근거 문장을 찾아 쓰세요")
    blue_box(doc, "💡 T/F 문제는 정답을 고른 후 반드시 근거 문장을 지문에서 찾아야 합니다!")
    for q_item in tf_qs:
        qp = p(doc, 5, 2)
        qp.paragraph_format.left_indent = Cm(0.2)
        run(qp, f"(     ) {q_item.get('no','')}.  ", bold=True, size=10)
        run(qp, q_item.get('statement',''), size=10)
        gp = p(doc, 1, 5)
        gp.paragraph_format.left_indent = Cm(1.5)
        run(gp, "근거 문장:  ", bold=True, size=9, color=C_GRAY)
        run(gp, "_" * 62, size=9, color=C_GRAY)

    blank_row(doc)
    sub_title(doc, "4-2.  내용 일치 5지선다  (2문제)")
    blue_box(doc, "⚠️ 모든 선택지를 지문과 꼼꼼히 대조하세요!")
    for key in ["content_choice_q1", "content_choice_q2"]:
        cq = content.get(key, {})
        if not cq:
            continue
        qp = p(doc, 6, 2)
        run(qp, cq.get('question',''), bold=True, size=10.5)
        add_choices(doc, cq.get('choices', []))

    # ────────────────────────────────────────────────────
    # SECTION 5: 유형별 문제
    # ────────────────────────────────────────────────────
    doc.add_page_break()
    section_title(doc, "🎯  SECTION 5 : 수능 유형별 문제")

    # A. 주제/요지/제목
    sub_type = content.get("subject_type", "주제")
    sub_title(doc, f"A.  {sub_type} 추론  (2문제)")
    blue_box(doc, f"💡 {sub_type}는 지문 전체를 관통하는 핵심입니다. 한 문단만 설명하는 선택지는 오답!")
    for key in ["subject_q1", "subject_q2"]:
        sq = content.get(key, {})
        if not sq:
            continue
        diff = sq.get('difficulty','')
        qp = p(doc, 8, 2)
        run(qp, f"【 {diff} 】  ", bold=True, size=10, color=C_MID)
        run(qp, sq.get('question',''), bold=True, size=10)
        add_choices(doc, sq.get('choices',[]))

    # B. 빈칸 추론
    blank_row(doc)
    sub_title(doc, "B.  빈칸 추론  (3문제  —  쉬움 → 어려움 순)")
    blue_box(doc, "💡 빈칸 전후 문장의 논리 흐름을 파악하세요. 빈칸만 읽으면 함정에 빠집니다!")
    for key in ["blank_q1", "blank_q2", "blank_q3"]:
        bq = content.get(key, {})
        if not bq:
            continue
        tp = p(doc, 8, 3)
        set_para_bg(tp, BG_YELLOW)
        run(tp, f"  [{bq.get('label','')}]  힌트: {bq.get('hint_kor','')}", bold=True, size=10, color=C_DARK)
        sp = p(doc, 3, 3)
        sp.paragraph_format.left_indent = Cm(0.5)
        sp.paragraph_format.line_spacing = Pt(20)
        run(sp, bq.get('passage_with_blank','').replace('______', '  [  ______  ]  '), size=10.5)
        add_choices(doc, bq.get('choices', []))

    # C. 어법 판단
    if include_grammar:
        blank_row(doc)
        doc.add_page_break()
        sub_title(doc, "C.  어법 판단  (5문제  —  맞으면 ○, 틀리면 ✕  후 수정)")
        blue_box(doc, "⚠️ 밑줄 친 부분만 보지 말고 문장 전체 구조를 분석하세요!")
        for item in content.get("grammar_items", []):
            gqp = p(doc, 7, 2)
            gqp.paragraph_format.left_indent = Cm(0.2)
            run(gqp, f"  {item.get('no','')}. [{item.get('point','')}]  ", bold=True, size=10, color=C_MID)
            sentence = item.get('sentence','')
            ul_part  = item.get('underline','')
            sp2 = p(doc, 2, 2)
            sp2.paragraph_format.left_indent = Cm(1.2)
            sp2.paragraph_format.line_spacing = Pt(19)
            if ul_part and ul_part in sentence:
                parts = sentence.split(ul_part, 1)
                run(sp2, parts[0], size=10.5)
                ur = run(sp2, ul_part, size=10.5, ul=True, bold=True)
                if not item.get('correct', True):
                    ur.font.color.rgb = C_RED
                run(sp2, parts[1], size=10.5)
            else:
                run(sp2, sentence, size=10.5)
            fp = p(doc, 2, 6)
            fp.paragraph_format.left_indent = Cm(1.5)
            run(fp, "판단: (  ○  /  ✕  )     수정: ", bold=True, size=9.5)
            run(fp, "_" * 25, size=9.5)

    # ────────────────────────────────────────────────────
    # SECTION 6: 서술형
    # ────────────────────────────────────────────────────
    if include_essay:
        doc.add_page_break()
        section_title(doc, "✍️  SECTION 6 : 서술형")

        sub_title(doc, "6-1.  단답형  —  영어로 답하세요")
        blue_box(doc, "💡 지문에서 근거를 찾아 그 표현을 활용해 쓰세요!")
        for q_item in content.get("short_answer_qs", []):
            qp2 = p(doc, 8, 2)
            run(qp2, f"  Q{q_item.get('no','')}.  ", bold=True, size=11, color=C_DARK)
            run(qp2, q_item.get('question',''), size=10.5)
            hp = p(doc, 2, 2)
            hp.paragraph_format.left_indent = Cm(0.8)
            run(hp, f"💬 힌트: {q_item.get('hint_kor','')}", size=9, color=C_GRAY, italic=True)
            for _ in range(3):
                lp = p(doc, 1, 1)
                lp.paragraph_format.left_indent = Cm(0.5)
                run(lp, "_" * 78, size=10)

        blank_row(doc)
        sub_title(doc, "6-2.  요약문 완성  —  빈칸에 알맞은 말을 쓰세요")
        blue_box(doc, "지문의 내용을 바탕으로 빈칸을 완성하세요.")
        tmpl = content.get("summary_template", "")
        if tmpl:
            smp = p(doc, 6, 6)
            smp.paragraph_format.left_indent = Cm(0.5)
            smp.paragraph_format.line_spacing = Pt(24)
            set_para_bg(smp, BG_LIGHT)
            run(smp, tmpl, size=10.5)

        blank_row(doc)
        sub_title(doc, "6-3.  핵심 문장 영작  —  한국어를 영어로 쓰세요 (괄호 안의 단어 활용)")
        for wt in content.get("writing_tasks", []):
            wp = p(doc, 8, 2)
            run(wp, f"  {wt.get('no','')}.  ", bold=True, size=11, color=C_MID)
            run(wp, wt.get('korean',''), size=10.5)
            hwp = p(doc, 1, 2)
            hwp.paragraph_format.left_indent = Cm(0.8)
            run(hwp, f"활용 어휘: {wt.get('hint_words','')}", size=9, color=C_GRAY, italic=True)
            for _ in range(3):
                llp = p(doc, 1, 1)
                llp.paragraph_format.left_indent = Cm(0.5)
                run(llp, "_" * 78, size=10)

    # ────────────────────────────────────────────────────
    # SECTION 7: 심화 학습
    # ────────────────────────────────────────────────────
    doc.add_page_break()
    section_title(doc, "💡  SECTION 7 : 심화 학습")

    sub_title(doc, "7-1.  나만의 문제 만들기")
    blue_box(doc, "이 지문을 바탕으로 직접 문제를 만들어 보세요. 출제자의 시각으로 지문을 보게 됩니다!")
    myq_tbl = doc.add_table(rows=7, cols=2)
    myq_tbl.style = 'Table Grid'
    for i, (lbl, val) in enumerate([
        ("유형 선택", "주제 / 요지 / 빈칸 / 내용일치 / 어법  중  →  (             )"),
        ("문제", ""), ("① 선택지", ""), ("② 선택지", ""),
        ("③ 선택지", ""), ("④ 선택지", ""), ("정답 & 해설", ""),
    ]):
        myq_tbl.rows[i].cells[0].text = lbl
        myq_tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        myq_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(myq_tbl.rows[i].cells[0], BG_LIGHT)
        myq_tbl.rows[i].cells[1].text = val
        myq_tbl.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        myq_tbl.rows[i].cells[0].width = Cm(2.8)
        myq_tbl.rows[i].cells[1].width = Cm(11.7)

    blank_row(doc)
    sub_title(doc, "7-2.  지문 한 줄 영어 요약")
    ols = content.get("one_line_summary", {})
    olp = p(doc, 6, 6)
    olp.paragraph_format.left_indent = Cm(0.5)
    set_para_bg(olp, BG_LIGHT)
    run(olp, "The passage mainly discusses  ___________________________________\n", size=11)
    run(olp, "because / by  ___________________________________________________.", size=11)

    blank_row(doc)
    sub_title(doc, "7-3.  오답 노트")
    err_tbl = doc.add_table(rows=5, cols=4)
    err_tbl.style = 'Table Grid'
    for i, h in enumerate(['문제 번호', '내가 고른 답', '정답', '틀린 이유 / 다음에 확인할 것']):
        cell = err_tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(cell, BG_GRAY)
    for row in err_tbl.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].add_run("\n").font.size = Pt(14)
    for i, w in enumerate([1.5, 2.0, 2.0, 9.0]):
        col_w(err_tbl, i, w)

    # ────────────────────────────────────────────────────
    # SECTION 8: 정답 및 해설
    # ────────────────────────────────────────────────────
    doc.add_page_break()
    section_title(doc, "📋  SECTION 8 : 정답 및 해설")
    blue_box(doc, "❗ 이 페이지는 문제를 모두 푼 후에 펼치세요!")

    # T/F 정답
    sub_title(doc, "■  T/F 정답")
    tf_rows = [(q.get('no',''), q.get('answer',''), q.get('evidence','')) for q in tf_qs]
    answer_table(doc, tf_rows, ['번호','정답','근거 문장'], [1.5, 1.5, 11.5])

    # 5지선다 정답 & 오답 해설
    blank_row(doc)
    sub_title(doc, "■  5지선다 정답 및 오답 해설")
    for label, key in [("내용일치 1", "content_choice_q1"), ("내용일치 2", "content_choice_q2"),
                       (f"{sub_type} 1 (표준)", "subject_q1"), (f"{sub_type} 2 (고난도)", "subject_q2")]:
        cq = content.get(key, {})
        if not cq:
            continue
        ap = p(doc, 6, 2)
        run(ap, f"  [{label}]  정답: ", bold=True, size=10.5, color=C_DARK)
        run(ap, cq.get('answer',''), bold=True, size=11, color=C_RED)
        exps = cq.get('explanations', {})
        exp_text = "\n".join(f"{'①②③④⑤'[int(k)-1] if k.isdigit() else k}  {v}" for k, v in exps.items())
        ep = p(doc, 2, 6)
        ep.paragraph_format.left_indent = Cm(0.5)
        ep.paragraph_format.line_spacing = Pt(20)
        set_para_bg(ep, BG_LIGHT)
        run(ep, exp_text, size=9.5)

    # 빈칸 정답
    blank_row(doc)
    sub_title(doc, "■  빈칸 추론 정답")
    blank_rows = []
    for key in ["blank_q1","blank_q2","blank_q3"]:
        bq = content.get(key, {})
        if bq:
            blank_rows.append((bq.get('label',''), bq.get('answer',''), bq.get('explanation','')))
    answer_table(doc, blank_rows, ['난이도','정답','근거'], [1.5, 4.0, 9.0])

    # 어법 정답
    if include_grammar:
        blank_row(doc)
        sub_title(doc, "■  어법 판단 정답")
        g_rows = []
        for item in content.get("grammar_items", []):
            judge = "○" if item.get('correct', True) else "✕"
            g_rows.append((item.get('no',''), judge, item.get('correction','—'), item.get('explanation','')))
        answer_table(doc, g_rows, ['번호','판단','수정','문법 해설'], [1.2, 1.2, 2.5, 9.6])

    # 서술형 모범 답안
    if include_essay:
        blank_row(doc)
        sub_title(doc, "■  서술형 모범 답안")
        essay_rows = []
        for q_item in content.get("short_answer_qs", []):
            essay_rows.append((f"Q{q_item.get('no','')}", q_item.get('model_answer','')))
        for i, ans in enumerate(content.get("summary_answers",[])):
            essay_rows.append((f"요약 ①②③④⑤"[i+2] if i < 5 else f"요약{i+1}", ans))
        for wt in content.get("writing_tasks", []):
            essay_rows.append((f"영작 {wt.get('no','')}", wt.get('model_answer','')))
        answer_table(doc, essay_rows, ['문항','모범 답안'], [2.0, 12.5])

    # ── 저장 → bytes ──────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
